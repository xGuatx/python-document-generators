from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ======================================================
# CONFIGURATION
# ======================================================
REPORT_INFO = {
    "title": "RAPPORT DE COMPÉTENCES",
    "subtitle": "Accompagnement et Projets Professionnels",
    "author": "Prénom NOM",
    "organization": "Organisme de Formation / Université [NOM]",
    "certification_name": "Titre / Certification [NOM]",
    "year": "2024–2025",
    "period_dates": "Du [DATE_DEBUT] au [DATE_FIN]",
    "company_name": "[NOM_ENTREPRISE]",
    "author_id": "prenom.nom"
}

OUTPUT_FILENAME = "Rapport_Stage_Template.docx"

# ======================================================
# DOCUMENT SETUP
# ======================================================
document = Document()

for section in document.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

style = document.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# ======================================================
# PAGE DE GARDE
# ======================================================
p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"{REPORT_INFO['title']}\n{REPORT_INFO['subtitle']}")
run.bold = True
run.font.size = Pt(18)

document.add_paragraph(f"\nNom / Prénom : {REPORT_INFO['author']}")
document.add_paragraph(f"Organisme de formation : {REPORT_INFO['organization']}")
document.add_paragraph(f"Diplôme / Certification : {REPORT_INFO['certification_name']}")
document.add_paragraph(f"Période et lieu : {REPORT_INFO['period_dates']} – {REPORT_INFO['company_name']}")
document.add_paragraph(f"Année : {REPORT_INFO['year']}")
document.add_paragraph(f"Auteur : {REPORT_INFO['author_id']}")

document.add_page_break()

# ======================================================
# SOMMAIRE
# ======================================================
document.add_heading("Sommaire", level=1)
sommaire = [
    "1. Introduction",
    "2. Contexte professionnel",
    "3. Présentation du projet / mission",
    "4. Analyse des besoins et objectifs",
    "5. Mise en œuvre des actions",
    "6. Évaluation et bilan",
    "7. Conclusion",
    "8. Annexes"
]
for item in sommaire:
    document.add_paragraph(item, style="List Number")

document.add_page_break()

# ======================================================
# 1. INTRODUCTION
# ======================================================
intro = (
    f"Dans le cadre de mon parcours en {REPORT_INFO['certification_name']}, j’ai réalisé une immersion professionnelle "
    f"{REPORT_INFO['period_dates']} au sein de {REPORT_INFO['company_name']}. "
    "Cette expérience avait pour but de [OBJECTIF_PRINCIPAL].\n\n"
    
    "Ce rapport présente [SUJET_DU_RAPPORT]. Il détaille les missions que j'ai menées, "
    "les compétences mobilisées et les résultats obtenus.\n\n"
)
document.add_heading("1. Introduction", level=1)
document.add_paragraph(intro)

document.add_page_break()

# ======================================================
# 2. CONTEXTE PROFESSIONNEL
# ======================================================
context = (
    "2.1 Présentation de l'organisme\n"
    f"{REPORT_INFO['company_name']} est une structure spécialisée dans [SECTEUR_ACTIVITE]. "
    "Ses missions principales sont [MISSIONS_PRINCIPALES].\n\n"
    
    "2.2 Présentation du service et de l’équipe\n"
    "J'ai évolué au sein de l'équipe [NOM_EQUIPE], composée de :\n"
    "- [ROLE_1]\n"
    "- [ROLE_2]\n\n"
    
    "2.3 Environnement de travail\n"
    "L'environnement se caractérise par [DESCRIPTION_ENVIRONNEMENT]."
)
document.add_heading("2. Contexte professionnel", level=1)
document.add_paragraph(context)

document.add_page_break()

# ======================================================
# 3. PRÉSENTATION DU PROJET / MISSION
# ======================================================
mission = (
    "3.1 Définition de la mission\n"
    "Ma mission principale consistait à [DESCRIPTION_MISSION].\n\n"
    
    "3.2 Enjeux\n"
    "Les enjeux identifiés étaient :\n"
    "- [ENJEU_1]\n"
    "- [ENJEU_2]\n\n"
    
    "3.3 Public ou Cible concernée\n"
    "[DESCRIPTION_CIBLE]"
)
document.add_heading("3. Présentation du projet / mission", level=1)
document.add_paragraph(mission)

document.add_page_break()

# ======================================================
# 4. ANALYSE ET OBJECTIFS
# ======================================================
analyse = (
    "4.1 Analyse de la situation\n"
    "L'état des lieux initial a révélé [OBSERVATIONS].\n\n"
    
    "4.2 Objectifs opérationnels\n"
    "Objectif 1 : [OBJECTIF_SMART_1]\n"
    "Objectif 2 : [OBJECTIF_SMART_2]"
)
document.add_heading("4. Analyse des besoins et objectifs", level=1)
document.add_paragraph(analyse)

document.add_page_break()

# ======================================================
# 5. MISE EN ŒUVRE
# ======================================================
actions = (
    "5.1 Méthodologie\n"
    "Pour atteindre ces objectifs, j'ai mis en place la méthodologie suivante : [METHODOLOGIE].\n\n"
    
    "5.2 Actions réalisées\n"
    "Action 1 : [DETAIL_ACTION_1]\n"
    "Action 2 : [DETAIL_ACTION_2]\n\n"
    
    "5.3 Moyens mobilisés\n"
    "- Humains : [MOYENS_HUMAINS]\n"
    "- Matériels / Techniques : [MOYENS_TECHNIQUES]"
)
document.add_heading("5. Mise en œuvre des actions", level=1)
document.add_paragraph(actions)

document.add_page_break()

# ======================================================
# 6. ÉVALUATION ET BILAN
# ======================================================
evaluation = (
    "6.1 Résultats obtenus\n"
    "[RESULTATS_QUANTITATIFS_QUALITATIFS]\n\n"
    
    "6.2 Écart par rapport aux objectifs\n"
    "[ANALYSE_ECARTS]\n\n"
    
    "6.3 Bilan personnel\n"
    "Cette expérience m'a permis de développer [COMPETENCES_ACQUISES]."
)
document.add_heading("6. Évaluation et bilan", level=1)
document.add_paragraph(evaluation)

document.add_page_break()

# ======================================================
# 7. CONCLUSION
# ======================================================
conclusion = (
    "En conclusion, ce projet au sein de {REPORT_INFO['company_name']} a été [QUALIFICATIF]. "
    "Il m'a permis de valider [ACQUIS]. Je suis désormais en mesure de [PERSPECTIVES]."
)
document.add_heading("7. Conclusion", level=1)
document.add_paragraph(conclusion)

# ======================================================
# 8. ANNEXES
# ======================================================
annexes = (
    "- Annexe 1 : [TITRE_ANNEXE_1]\n"
    "- Annexe 2 : [TITRE_ANNEXE_2]"
)
document.add_heading("8. Annexes", level=1)
document.add_paragraph(annexes)

# ======================================================
# ENREGISTREMENT
# ======================================================
document.save(OUTPUT_FILENAME)
print(f"Document généré : {OUTPUT_FILENAME}")
