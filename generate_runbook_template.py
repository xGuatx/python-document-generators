#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ========= Données Configurable =========
COMPANY_NAME = "Entreprise XYZ"
RUNBOOK = {
    "titre": f"Runbook 1-Page — {COMPANY_NAME}",
    "but_cibles": [
        "Maintien de l'activité critique",
        "RTO cible = 4h / RPO = 1h",
        "Assurer la continuité de service"
    ],
    "declencheurs": [
        "Alerte monitoring critique",
        "Panne infra majeure",
        "Perte de connectivité > X%"
    ],
    "objectifs_immediats": [
        "Sécuriser les données",
        "Restaurer les services prioritaires",
        "Communiquer avec les parties prenantes"
    ],
    "etapes": [
        ("Alerter", "Responsable Incident [Nom / Tel]"),
        ("Isoler", "Isoler les composants affectés"),
        ("Basculer", "Activer le plan de bascule (Failover)"),
        ("Vérifier", "Vérifier l'intégrité et la disponibilité"),
        ("Communiquer", "Envoyer statut aux équipes et clients"),
        ("Clôture", "Valider le retour à la normale")
    ],
    "contacts": [
        ("Responsable Crise", "[Nom]", "+33 6 XX XX XX XX", "Tel / Chat"),
        ("Responsable Tech", "[Nom]", "+33 6 XX XX XX XX", "Tel / Chat"),
        ("Support Vendor", "[Nom]", "+33 X XX XX XX XX", "Portal"),
        ("Communication", "[Nom]", "+33 6 XX XX XX XX", "Email"),
    ],
    "temps": ["Alerte : 5’", "Analyse : 15’", "Action : 60’", "Validation : continu"],
    "planB": "Activation procédure manuelle dégradée",
    "pied": f"© {COMPANY_NAME} — Version 1.0 — Confidentiel"
}

CONSIGNES = {
    "titre": "Consignes Runbook",
    "objectifs_session": [
        "Relire les procédures",
        "Valider les contacts",
        "Tester le scénario"
    ],
    "checklist": [
        "Déclencheur identifié",
        "Contacts à jour",
        "Outils accessibles"
    ]
}

# ========= Mise en forme =========
def set_margins(section, top=1.5, bottom=1.5, left=2.0, right=2.0):
    section.top_margin = Cm(top); section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left); section.right_margin = Cm(right)

def title(doc, txt, size=20, color=(20,40,80)):
    p = doc.add_paragraph(); r = p.add_run(txt)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = RGBColor(*color)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT; return p

def h(doc, txt):
    p = doc.add_paragraph(); r = p.add_run(txt); r.bold = True; r.font.size = Pt(13); return p

def bullets(doc, items):
    for x in items:
        p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(0); p.add_run(x)

def numbered_steps(doc, steps):
    for (verb, desc) in steps:
        p = doc.add_paragraph(style='List Number'); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{verb} — "); r.bold = True
        p.add_run(desc)

def contacts_table(doc, rows):
    tbl = doc.add_table(rows=1, cols=4); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, t in enumerate(["Rôle", "Nom", "Téléphone", "Accès"]):
        cell = tbl.rows[0].cells[i]; r = cell.paragraphs[0].add_run(t); r.bold = True
    for role, nom, tel, acc in rows:
        row = tbl.add_row().cells; row[0].text=role; row[1].text=nom; row[2].text=tel; row[3].text=acc

def checkboxes(doc, items):
    for it in items:
        p = doc.add_paragraph(); p.add_run("☑ ").bold = True; p.add_run(it)

def small_footer(doc, txt):
    p = doc.add_paragraph(); r = p.add_run(txt); r.font.size = Pt(8); r.italic = True

def fiche_revue(doc, idx=1):
    h(doc, f"Fiche de revue n°{idx}")
    p = doc.add_paragraph(); p.add_run("Binôme : __________________   Date : ____/____/______")
    doc.add_paragraph()
    tbl = doc.add_table(rows=4, cols=3); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Critère", "Score", "Commentaire"]
    for i, t in enumerate(headers):
        r = tbl.rows[0].cells[i].paragraphs[0].add_run(t); r.bold = True
    doc.add_paragraph()

# ========= Génération =========
def build_doc(path="Runbook_Template.docx"):
    doc = Document(); set_margins(doc.sections[0])

    # Page 1 — Runbook
    title(doc, RUNBOOK["titre"])
    h(doc, "But & cibles"); bullets(doc, RUNBOOK["but_cibles"])
    h(doc, "Déclencheur"); bullets(doc, RUNBOOK["declencheurs"])
    h(doc, "Objectifs immédiats"); bullets(doc, RUNBOOK["objectifs_immediats"])
    h(doc, "Étapes"); numbered_steps(doc, RUNBOOK["etapes"])
    h(doc, "Contacts"); contacts_table(doc, RUNBOOK["contacts"])
    h(doc, "Temps estimés"); bullets(doc, RUNBOOK["temps"])
    h(doc, "Plan B"); doc.add_paragraph(RUNBOOK["planB"])
    small_footer(doc, RUNBOOK["pied"])

    doc.add_page_break()

    # Page 2 — Consignes
    title(doc, CONSIGNES["titre"])
    h(doc, "Objectifs"); bullets(doc, CONSIGNES["objectifs_session"])
    h(doc, "Checklist"); checkboxes(doc, CONSIGNES["checklist"])

    doc.save(path); print(f"✅ Document généré : {path}")

if __name__ == "__main__":
    build_doc()
